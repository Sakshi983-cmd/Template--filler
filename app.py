import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import tempfile
import os

# 🎯 Page setup
st.set_page_config(page_title="Task 3 - Insurance Template Filler", layout="centered")
st.title("📄 Insurance Template Auto-Filler using PDF + Mock LLM")

# 📤 Upload Inputs
template_file = st.file_uploader("Upload Insurance Template (.docx)", type="docx")
pdf_files = st.file_uploader("Upload PDF Photo Reports", type="pdf", accept_multiple_files=True)

# 🧠 Extract text from PDFs
def extract_text_from_pdfs(files):
    text = ""
    for file in files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        doc = fitz.open(tmp_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    return text

# 🤖 Mock LLM to fill template (no API needed)
def fill_template_with_llm(template_text, pdf_text):
    return f"""📋 Auto-Filled Insurance Report

PDF Extract Preview:
{pdf_text[:300]}...

Identified Issues:
- Roof leakage
- Ceiling staining
- Wall cracks and paint damage

Recommendation:
Immediate repair recommended. Insurance claim approved for structural damage.
"""

# 🔄 Main Execution
if st.button("Generate Filled Template") and template_file and pdf_files:
    st.info("📤 Extracting text from PDFs...")
    pdf_text = extract_text_from_pdfs(pdf_files)

    with st.expander("📄 Preview Extracted PDF Text"):
        st.text_area("Extracted Text", pdf_text[:2000], height=200)

    st.info("📄 Reading insurance template...")
    doc = Document(template_file)
    template_text = "\n".join([p.text for p in doc.paragraphs])

    st.info("🧠 Mocking LLM response...")
    filled_output = fill_template_with_llm(template_text, pdf_text)

    st.success("✅ Template filled successfully!")

    filled_doc = Document()
    for line in filled_output.split("\n"):
        filled_doc.add_paragraph(line.strip())

    output_path = os.path.join(tempfile.gettempdir(), "filled_template.docx")
    filled_doc.save(output_path)

    with open(output_path, "rb") as f:
        st.download_button("⬇️ Download Filled Template", f, file_name="filled_template.docx")
